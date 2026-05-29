"""render_docx — convert a structured Markdown draft to HSE-style .docx.

Used by grant_application_drafter (and other drafters) to produce a portal-
ready .docx alongside the canonical .md.

Style choices (Russian academic standard):
  - Page size A4, margins 2cm
  - Font: Times New Roman 12pt, single line spacing
  - Headings: bold, 14pt for H1, 13pt for H2, 12pt for H3
  - Tables: visible borders, centered headers
  - Page numbers in footer (bottom-right)

Markdown features handled:
  - # / ## / ### headings → Word heading levels with proper styling
  - **bold** and *italic* runs
  - - bullet lists, 1. numbered lists
  - | table | rows | with header row
  - HTML comments <!-- ... --> stripped (drafter metadata)
  - blank lines → paragraph breaks

Usage:
  render_docx.markdown_to_docx(md_text, "/path/to/out.docx", title="Заявка ...")
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


# ────────────────────────────────────────────────────────────────────
# Style helpers
# ────────────────────────────────────────────────────────────────────
def _set_default_style(doc: Document) -> None:
    """Set body font + page geometry to Russian-academic standard."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # East-Asian font setting needed for Cyrillic in some viewers
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)  # left wider — РНФ template convention
        section.right_margin = Cm(1.5)


def _add_page_numbers(doc: Document) -> None:
    """Insert page number in footer (bottom-right)."""
    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = footer_para.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)


# ────────────────────────────────────────────────────────────────────
# Inline markdown → runs
# ────────────────────────────────────────────────────────────────────
INLINE_PAT = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Split text into bold/italic/code runs and add to paragraph."""
    pos = 0
    for m in INLINE_PAT.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(11)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ────────────────────────────────────────────────────────────────────
# Block parser
# ────────────────────────────────────────────────────────────────────
def _strip_html_comments(md: str) -> str:
    return re.sub(r"<!--.*?-->", "", md, flags=re.DOTALL)


def _parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Read a Markdown table starting at lines[start]. Return (rows, end_index)."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # Skip the separator row (---|---|---)
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            cell.text = row[ci] if ci < len(row) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if ri == 0:
                        run.bold = True


# ────────────────────────────────────────────────────────────────────
# Main entry point
# ────────────────────────────────────────────────────────────────────
def markdown_to_docx(md_text: str, out_path: str | Path, title: str = "") -> str:
    doc = Document()
    _set_default_style(doc)
    _add_page_numbers(doc)

    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in h.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
            run.bold = True

    md = _strip_html_comments(md_text)
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Headings
        m_h = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m_h:
            level = min(len(m_h.group(1)), 4)
            text = m_h.group(2).strip()
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = "Times New Roman"
                size = {1: 14, 2: 13, 3: 12, 4: 12}[level]
                run.font.size = Pt(size)
                run.bold = True
            i += 1
            continue

        # Horizontal rule — skip
        if re.match(r"^-{3,}$", line.strip()) or re.match(r"^={3,}$", line.strip()):
            i += 1
            continue

        # Table block
        if line.strip().startswith("|"):
            rows, end_i = _parse_table_block(lines, i)
            _add_table(doc, rows)
            i = end_i
            continue

        # Bullet list (possibly multi-line)
        if re.match(r"^[-*•]\s+", line.strip()):
            while i < len(lines) and re.match(r"^[-*•]\s+", lines[i].strip()):
                item_text = re.sub(r"^[-*•]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, item_text)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line.strip()):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item_text = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, item_text)
                i += 1
            continue

        # Regular paragraph (may span multiple lines until blank)
        para_lines = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(
            ("#", "|", "-", "*", "•")
        ) and not re.match(r"^\d+\.\s+", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return str(out_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 3:
        print("usage: render_docx.py <input.md> <output.docx> [title]")
        sys.exit(1)
    md = Path(sys.argv[1]).read_text(encoding="utf-8")
    title = sys.argv[3] if len(sys.argv) > 3 else ""
    out = markdown_to_docx(md, sys.argv[2], title)
    print(f"✓ wrote {out}")
